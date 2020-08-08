import docx
import os
import subprocess
import time
import datetime
import sys
import re

clr = "\033[1;36;40m"
clear = "\033[0;00;00m"

print(clr + "PCI_parse" + clear + " -PCI Segmentation Testing Table Generator -")
print("")
error = 0


gnmap_file = sys.argv[-1]
fexist = 0

try:
	f = open(gnmap_file)
	f.close()
	fexist = 4
except FileNotFoundError:
	print('File does not exist')
	error = 1
	trigger = 1
	fexist = 0
	os.system('exit')


try:
	random = sys.argv[1]	

except IndexError:
	print("Usage: python3 seg_test.py <filename>.gnmap")
	os.system('exit')
	error = 1


trigger = 1
try:
	if os.path.splitext(sys.argv[1])[1] == ".gnmap":
		if fexist == 4:
			print("[+] Generating Report")
		else:
			print(" ")

	else:
		print("Usage: python3 seg_test.py <filename>.gnmap")
		os.system('exit')
		error = 1
		trigger = 1

except IndexError:
	if trigger == 1:
		print(" ")
	else:	
		print("Usage: python3 seg_test.py <filename>.gnmap")
		os.system('exit')
		error = 1
	



doc = docx.Document()
doc.add_heading('Network Segmentation Testing', level=1)
doc.add_paragraph('Following PCI in-scope hosts and their ports are accessible from PCI out of scope hosts')

os.system("./grep_ip.sh" + " " + gnmap_file + " " + "2>/dev/null")
os.system("./grep_ports.sh" + " " + gnmap_file + " " + "2>/dev/null")

##############################################
f1 = open('data/ip.txt', 'r')
raw_ip_address=[]

for ip in f1:
    raw_ip_address.append(ip)

ip_address = []
for element_1 in raw_ip_address:
    ip_address.append(element_1.strip())

f2 = open('data/ports.txt', 'r')
raw_ports=[]
ports=[]

for port in f2:
    raw_ports.append(port)

for element_2 in raw_ports:
    ports.append(element_2.strip())


'''
i = 0
for i < ip_len:
    element_combine.append(ip_address[i],ports[i])
    i += 1
'''

ip_len = (len(ip_address))
ports_len = (len(ports))

element_combine = dict(zip(ip_address, map(str, ports)))

#######################################################


menuTable = doc.add_table(rows=1,cols=2)
menuTable.style= 'Table Grid'
hdr_Cells = menuTable.rows[0].cells
hdr_Cells[0].text = 'Hosts'
hdr_Cells[1].text = 'Ports'


for Hosts, Ports in element_combine.items():
    row_Cells = menuTable.add_row().cells
    row_Cells[0].text= str(Hosts)
    row_Cells[1].text = str(Ports)

if error == 0:
	doc.save('Network-PCI-Segmentation-Testing-Report.docx')
else:
	error = 2